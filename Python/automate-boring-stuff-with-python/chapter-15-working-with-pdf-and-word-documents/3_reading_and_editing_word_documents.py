# The Python-Docx third-party module can read and write .docx Word files.
# Open a Word file with docx.Document()
# Access one of the Paragraph objects from the "paragraphs" member variable, which is a list of Paragraph objects.
# Paragraph objects have a "text" member variable containing the text as a string value.
# Paragraphs are composed of "runs".  The "runs" member variable of a Paragraph object contains a list of Run objects.
# Run objects also have a "text" member variable.
# Run objects have a "bold", "italic", and "underline" member variables which can be set to True or False.
# Paragraph and run objects have a "style" member variable that can be set to one of Word's built-in styles.
# Word files can be created by calling add_paragraph() and add_run() to append text content.

# install docx module: pip3 install python-docx

import docx, os


os.chdir(path=os.path.dirname(__file__))


doc = docx.Document("demo.docx")  # open a Word file

print(doc.paragraphs)  # list of all paragraphs

print(doc.paragraphs[0].text)  # text of first paragraph

print(doc.paragraphs[1].text)  # text of second paragraph

p = doc.paragraphs[1]  # second paragraph

print(p.runs)  # list of runs in paragraph

print(p.runs[1].text)  # text of second run

print(p.runs[2].bold)  # True if run is bold

print(p.runs[1].bold == None)  # True if run is not bold

print(p.runs[4].italic)  # True if run is italic

# change text of fourth run
p.runs[4].text = "italic and underlined."

doc.save("demo2.docx")  # save changes to Word file


print(p.style)  # style of paragraph
p.style = "Title"  # set paragraph style to "Title"
doc.save("demo3.docx")  # save changes to Word file


# create a blank new document
d = docx.Document()

# add new paragraph
d.add_paragraph("This is a new paragraph.")

d.add_paragraph("This is another new paragraph.")

d.save("demo4.docx")  # save changes to Word file

p = d.paragraphs[0]  # first paragraph
p.add_run("This is a new run.")  # add new run to paragraph
print(p.runs)
p.runs[1].bold = True  # make second run bold
d.save("demo5.docx")  # save changes to Word file